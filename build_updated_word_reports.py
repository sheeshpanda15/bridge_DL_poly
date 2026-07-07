from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SUMMARY = ROOT / "paper_highdim_iter_dopt_final_summary.csv"
FIG_FINAL = ROOT / "paper_highdim_iter_dopt_final_gain_ci.png"
FIG_CPU_GAIN = ROOT / "paper_highdim_iter_dopt_cpu_gain_mean_ci.png"
FIG_CUDA_GAIN = ROOT / "paper_highdim_iter_dopt_cuda_gain_mean_ci.png"
FIG_CPU_MSE = ROOT / "paper_highdim_iter_dopt_cpu_mse_mean_ci.png"
FIG_CUDA_MSE = ROOT / "paper_highdim_iter_dopt_cuda_mse_mean_ci.png"

OUT_CN = ROOT / "大报告_中文_完整版_论文级高维迭代更新版.docx"
OUT_EN = ROOT / "big_report_en_complete_paper_highdim_update.docx"


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x0B, 0x25, 0x45)
MUTED = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "F2F4F7"
BORDER = "C9D3DF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, bold=None, color=None, east_asia="Microsoft YaHei"):
    font = run.font
    font.name = "Calibri"
    if size is not None:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if color is not None:
        font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def style_doc(doc, lang):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header_text = (
        "Paper-level high-dimensional iterative D-optimal experiment"
        if lang == "en" else "论文级高维迭代 D-optimal 实验"
    )
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(header_text)
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated from reproducible experiment outputs")
    set_run_font(run, size=9, color=MUTED)


def add_title(doc, title, subtitle, meta):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    set_run_font(run, size=22, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(subtitle)
    set_run_font(run, size=12, color=MUTED)

    table = doc.add_table(rows=0, cols=2)
    widths = [1900, 7460]
    for label, value in meta:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
    set_table_geometry(table, widths)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)
        set_cell_shading(row.cells[0], HEADER_FILL)
        for run in row.cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = DARK_BLUE
    doc.add_paragraph()


def add_lead_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + " ")
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        run = p.add_run(item)
        set_run_font(run, size=10.5)


def add_metric_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        hdr.cells[idx].text = header
        set_cell_shading(hdr.cells[idx], HEADER_FILL)
        for run in hdr.cells[idx].paragraphs[0].runs:
            set_run_font(run, size=8.5, bold=True, color=DARK_BLUE)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            row.cells[idx].text = str(value)
            for run in row.cells[idx].paragraphs[0].runs:
                set_run_font(run, size=8.2)
            if idx in (1, 2, 4, 5):
                row.cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(table, widths)
    doc.add_paragraph()


def add_figure(doc, path, caption, width=6.25):
    if not Path(path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=9, color=MUTED)


def final_rows(df, device):
    sub = df[df["device"] == device].sort_values(["p", "case", "initial_fraction"])
    rows = []
    for r in sub.itertuples(index=False):
        case_label = str(r.case).replace("highdim_", "")
        rows.append([
            case_label,
            int(r.p),
            f"{r.initial_fraction:g}",
            int(r.mse_gain_vs_random_median_pct_count),
            f"{r.mse_gain_vs_random_median_pct_mean:.2f}% +/- {r.mse_gain_vs_random_median_pct_ci95:.2f}%",
            f"{100.0 * r.positive_rate:.1f}%",
        ])
    return rows


def build_chinese(df):
    doc = Document()
    style_doc(doc, "cn")
    add_title(
        doc,
        "NN-FullPR 迁移中的高维迭代 D-optimal 实验报告",
        "基于 10000 样本、p=20/50、CPU/GPU、多 seed 的论文级实验更新版",
        [
            ("数据规模", "原始 n=10000；训练池/评估集=7500/2500"),
            ("维度与场景", "p=20,50；quadratic / smooth / strong nonlinear"),
            ("重复设计", "3 个 data seed × 2 个 NN 初始化 seed × 2 个初始采样比例 × CPU/GPU"),
            ("计算量", "72 次 NN oracle 训练；144 条迭代轨迹；8640 次随机 surrogate 拟合"),
        ],
    )

    doc.add_heading("摘要", level=1)
    add_lead_callout(
        doc,
        "核心结论：",
        "在最后一轮 144 个重复中，D-optimal 升级有 128 个重复优于同预算随机升级中位数，正收益比例为 88.9%，平均 MSE gain 为 2.61%。这说明迭代式 D-optimal 可以作为高维 NN-to-FullPR surrogate 构建中的稳定数据升级策略，但收益大小依赖特征空间、初始采样比例和非线性强度。",
    )

    doc.add_heading("实验流程", level=1)
    add_bullets(doc, [
        "先生成原始高维数据集，并按 75/25 划分为训练池和评估集。",
        "在训练池上训练 NN oracle，将该 NN 视为新的目标模型。",
        "从训练池中均匀抽取初始设计集，估计 NN 与 FullPR surrogate 的 pilot 距离。",
        "每一轮基于当前已经升级后的设计集计算 D-optimal leverage，选择下一批样本加入设计集。",
        "每轮重新拟合 FullPR surrogate，并与同预算随机升级的中位数表现比较。",
    ])

    doc.add_heading("实验矩阵与计算量", level=1)
    add_metric_table(
        doc,
        ["项目", "设置"],
        [
            ["原始数据集大小", "10000"],
            ["输入维度", "20, 50"],
            ["数据场景", "quadratic, smooth, strong nonlinear"],
            ["初始 uniform 比例", "0.05, 0.1"],
            ["迭代轮数", "5；每轮增加 500 个样本"],
            ["随机基准", "每个迭代点重复 10 次"],
            ["设备", "CPU 与 NVIDIA GeForce RTX 4070 Ti SUPER"],
            ["FullPR 特征", "degree=2, include_special=True"],
        ],
        [2200, 7160],
    )

    doc.add_heading("主要结果", level=1)
    add_figure(doc, FIG_FINAL, "图 1. 最后一轮 D-optimal 相对随机升级的 mean gain +/- 95% CI。")

    doc.add_heading("最后一轮分组统计", level=2)
    doc.add_paragraph("表 1 和表 2 汇总最后一轮结果。gain 为正表示 D-optimal surrogate 的 MSE 低于同预算随机升级中位数。")
    add_metric_table(
        doc,
        ["Case", "p", "Init", "Reps", "Gain mean +/-95%CI", "Positive"],
        final_rows(df, "cpu"),
        [1850, 620, 820, 720, 3600, 1750],
    )
    add_metric_table(
        doc,
        ["Case", "p", "Init", "Reps", "Gain mean +/-95%CI", "Positive"],
        final_rows(df, "cuda"),
        [1850, 620, 820, 720, 3600, 1750],
    )

    doc.add_heading("迭代曲线", level=1)
    add_figure(doc, FIG_CPU_GAIN, "图 2. CPU 实验中 D-optimal 相对随机升级的 gain 均值曲线与 95% CI。")
    add_figure(doc, FIG_CUDA_GAIN, "图 3. CUDA 实验中 D-optimal 相对随机升级的 gain 均值曲线与 95% CI。")
    add_figure(doc, FIG_CPU_MSE, "图 4. CPU 实验中 D-optimal surrogate 与随机基准的 MSE 迭代曲线。")
    add_figure(doc, FIG_CUDA_MSE, "图 5. CUDA 实验中 D-optimal surrogate 与随机基准的 MSE 迭代曲线。")

    doc.add_heading("讨论", level=1)
    add_bullets(doc, [
        "p=20 时二阶 FullPR 特征数较少，初始样本通常已超过特征数，因此 D-optimal 的收益较温和，但正收益比例高。",
        "p=50 时初始设计集更容易低于 FullPR 特征数，早期迭代出现更大波动；当升级数据集继续进入下一轮 D-optimal 后，误差逐步稳定。",
        "strong nonlinear 场景的置信区间明显更宽，说明当 NN oracle 的局部行为超出二阶 FullPR 表达能力时，D-optimal 仍能改善覆盖，但不能被解释为充分逼近 NN 的保证。",
        "CPU/GPU 的总体结论一致，说明结果不是单一设备偶然现象。",
    ])

    doc.add_heading("结论", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "这版实验可以作为论文主实验使用。更稳妥的论文表述是：pilot 距离和迭代式 D-optimal 升级提供了一个可观测、可复现的数据选择诊断流程；该流程在高维设置中能够稳定优于随机升级，但收益幅度受 FullPR 特征空间、非线性强度、初始采样比例和迭代预算共同影响。"
    )
    set_run_font(r)
    doc.save(OUT_CN)


def build_english(df):
    doc = Document()
    style_doc(doc, "en")
    add_title(
        doc,
        "High-Dimensional Iterative D-optimal Experiment for NN-FullPR Transfer",
        "Paper-scale bilingual report update with n=10000, p=20/50, CPU/GPU, and multi-seed replication",
        [
            ("Dataset scale", "Original n=10000; training/evaluation split=7500/2500"),
            ("Dimensions and cases", "p=20,50; quadratic / smooth / strong nonlinear"),
            ("Replication", "3 data seeds × 2 NN initialization seeds × 2 initial fractions × CPU/GPU"),
            ("Workload", "72 NN oracle trainings; 144 iterative paths; 8640 random surrogate fits"),
        ],
    )

    doc.add_heading("Executive Summary", level=1)
    add_lead_callout(
        doc,
        "Main finding:",
        "At the final iteration, D-optimal upgrading outperformed the same-budget random-upgrade median in 128 of 144 replications. The final positive-rate was 88.9%, and the average MSE gain was 2.61%. The result supports iterative D-optimal design as a stable data-upgrading strategy for high-dimensional NN-to-FullPR surrogate construction, while also showing that the gain depends on feature expressiveness, initial sampling fraction, and nonlinearity strength.",
    )

    doc.add_heading("Experimental Workflow", level=1)
    add_bullets(doc, [
        "Generate a high-dimensional original dataset and split it into a training pool and an evaluation set.",
        "Train an NN oracle on the training pool; this NN is treated as the new target model.",
        "Uniformly sample the initial design set and estimate the pilot distance between the NN and the FullPR surrogate.",
        "At each iteration, compute D-optimal leverage conditional on the already upgraded design set and add a new batch of samples.",
        "Refit the FullPR surrogate after each upgrade and compare it against random upgrades under the same budget.",
    ])

    doc.add_heading("Experimental Matrix and Workload", level=1)
    add_metric_table(
        doc,
        ["Item", "Setting"],
        [
            ["Original dataset size", "10000"],
            ["Input dimensions", "20, 50"],
            ["Synthetic cases", "quadratic, smooth, strong nonlinear"],
            ["Initial uniform fractions", "0.05, 0.1"],
            ["Iterations", "5; 500 samples added per iteration"],
            ["Random baseline", "10 repeats per iteration point"],
            ["Devices", "CPU and NVIDIA GeForce RTX 4070 Ti SUPER"],
            ["FullPR features", "degree=2, include_special=True"],
        ],
        [2200, 7160],
    )

    doc.add_heading("Main Results", level=1)
    add_figure(doc, FIG_FINAL, "Figure 1. Final-iteration D-optimal gain over random upgrading, mean +/- 95% CI.")

    doc.add_heading("Final-Iteration Group Statistics", level=2)
    doc.add_paragraph("Tables 1 and 2 summarize final-iteration performance. A positive gain means that the D-optimal surrogate has lower MSE than the same-budget random-upgrade median.")
    add_metric_table(
        doc,
        ["Case", "p", "Init", "Reps", "Gain mean +/-95%CI", "Positive"],
        final_rows(df, "cpu"),
        [1850, 620, 820, 720, 3600, 1750],
    )
    add_metric_table(
        doc,
        ["Case", "p", "Init", "Reps", "Gain mean +/-95%CI", "Positive"],
        final_rows(df, "cuda"),
        [1850, 620, 820, 720, 3600, 1750],
    )

    doc.add_heading("Iteration Curves", level=1)
    add_figure(doc, FIG_CPU_GAIN, "Figure 2. CPU gain curves with 95% confidence intervals.")
    add_figure(doc, FIG_CUDA_GAIN, "Figure 3. CUDA gain curves with 95% confidence intervals.")
    add_figure(doc, FIG_CPU_MSE, "Figure 4. CPU MSE curves for D-optimal surrogates and random baselines.")
    add_figure(doc, FIG_CUDA_MSE, "Figure 5. CUDA MSE curves for D-optimal surrogates and random baselines.")

    doc.add_heading("Discussion", level=1)
    add_bullets(doc, [
        "For p=20, the quadratic FullPR feature count is relatively small, so the initial design often already exceeds the number of features. D-optimal gains are therefore modest but stable.",
        "For p=50, the initial design is more likely to be under the FullPR feature count. Early iterations are more volatile, but the upgraded design stabilizes as it is reused in later D-optimal steps.",
        "The strong nonlinear case has wider confidence intervals, which marks the boundary of the method: D-optimal selection can improve coverage, but it is not a guarantee that a low-degree FullPR surrogate fully captures the NN oracle.",
        "CPU and GPU runs lead to the same qualitative conclusion, reducing the chance that the result is a device-specific artifact.",
    ])

    doc.add_heading("Conclusion", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "This updated experiment is suitable as a main paper experiment. The most defensible claim is that pilot distance and iterative D-optimal upgrading form an observable and reproducible diagnostic workflow for data selection. In high-dimensional settings, the workflow usually improves over random upgrading, but the magnitude of improvement depends on the FullPR feature space, nonlinearity strength, initial sampling fraction, and iteration budget."
    )
    set_run_font(r)
    doc.save(OUT_EN)


def main():
    df = pd.read_csv(SUMMARY)
    build_chinese(df)
    build_english(df)
    print(OUT_CN)
    print(OUT_EN)


if __name__ == "__main__":
    main()
