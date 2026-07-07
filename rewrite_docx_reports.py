from copy import deepcopy
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent


def add_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    p.style = style or paragraph.style
    if text:
        p.add_run(text)
    return p


def add_after_element(element, parent, text="", style="Normal"):
    new_p = OxmlElement("w:p")
    element.addnext(new_p)
    p = Paragraph(new_p, parent)
    p.style = style
    if text:
        p.add_run(text)
    return p


def add_table_after(paragraph, rows, cols):
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Inches(6.2))
    paragraph._element.addnext(table._element)
    return table


def set_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def find_para(doc, exact=None, contains=None, style=None):
    for p in doc.paragraphs:
        t = p.text.strip()
        if exact is not None and t == exact:
            return p
        if contains is not None and contains in t:
            if style is None or p.style.name == style:
                return p
    raise ValueError(f"paragraph not found: exact={exact!r}, contains={contains!r}")


def add_bullets_after(anchor, bullets, style="List Bullet"):
    p = anchor
    for text in bullets:
        p = add_after(p, text, style)
    return p


def add_caption_after(anchor, text, image_path):
    cap = add_after(anchor, text, "Caption")
    pic = add_after(cap, "", "Normal")
    pic.alignment = 1
    pic.add_run().add_picture(str(ROOT / image_path), width=Inches(6.2))
    return pic


def fill_table(table, headers, rows):
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
    try:
        table.style = "Table Grid"
    except Exception:
        pass


def pct(x):
    return f"{100 * float(x):.1f}%"


def imp(x):
    return f"{float(x):+.1f}%"


def baseline_rows(case_name=None):
    df = pd.read_csv(ROOT / "results_clean_nn_baseline_summary.csv")
    if case_name is not None:
        df = df[df["case"] == case_name]
    return df


def chinese_report():
    doc = Document(str(ROOT / "大报告_中文_完整版.docx"))

    set_text(find_para(doc, contains="本报告整合了：数学基础"),
             "本报告整合了：数学基础、方法论、全部实验图表的逐子图解读、3000 次重复模拟的定量相关性结果、以 NN 为基准的模型胜负分析，以及可直接用于论文的结论。")
    set_text(find_para(doc, contains="本研究回答一个根本问题"),
             "本研究回答一个根本问题：神经网络是否在“偷偷地”做多项式回归？我们用“形状距离”这把几何尺子，把这个模糊直觉变成可测量的科学命题，并在 5 个数据集、3 种激活函数、3 种网络深度、共 3000 次重复模拟上检验。新增的 NN 基准胜负图进一步回答“谁更准”：不是所有多项式都比 NN 好，FullPR 只在部分结构较清晰的数据集上胜出。核心发现如下：")
    set_text(doc.paragraphs[10],
             "核心命题成立且普适：神经网络与多项式回归的几何接近度，强烈预测二者的性能接近度（清洗后全体 Spearman ρ = 0.839，p ≈ 0；各数据集 ρ 介于 0.440–0.870，均显著）。")
    set_text(doc.paragraphs[13],
             "逐层泰勒展开（LayerTaylorPR）稀有但灾难性失效：绝大多数情形下能复刻 NN，但一旦突触电位冲出收敛域，MSE 会爆炸；在“以 NN 为基准”的预测胜负中，它总体胜率仅 5.6%，因此更适合作为解释/复刻工具，而不是稳健预测替代品。")
    add_bullets_after(doc.paragraphs[13], [
        "以 NN 为基准的预测胜负并非单向结论：FullPR 在 paper_poly3（胜率 59.5%，中位改善 +62.6%）和 smooth_nonlinear（胜率 58.5%，中位改善 +37.0%）中更常胜出；但在 paper_poly4（胜率 36.7%，中位改善 -97.0%）和 smooth_nonlinear_rand（胜率 14.4%，中位改善 -169.1%）中，NN 明显更稳。",
    ], "List Number")

    set_text(find_para(doc, contains="变体：NN_mse_vs_y"),
             "变体：NN_mse_vs_y（NN vs 真值）、FPR_mse_vs_y（FullPR vs 真值）、FPR_mse_vs_NN（FullPR vs NN 输出）、LTPR_mse_vs_NN（LTPR vs NN 输出）。后两者以 NN 输出为“真值”，因为关心的是“像不像 NN”。新增的 NN 基准指标包括：*_mse_delta_vs_NN（模型 MSE 减 NN MSE，负值表示模型更准）、*_improve_vs_NN_pct（相对 NN 的改善百分比，正值表示模型更好）、*_better_than_NN（该 run 是否击败 NN）。")
    set_text(find_para(doc, contains="子图④　NN MSE vs FullPR MSE"),
             "子图④　NN MSE vs FullPR MSE（y=x 线）：点在线下=FullPR 更准，线上=NN 更准。新加入的 nn_baseline 图把这个判断汇总为胜率和中位改善，更适合报告“哪个更好”。")

    anchor = find_para(doc, exact="第四部分　定量发现")
    p = add_after(anchor, "4.0　以 NN 为基准的模型胜负", "Heading 2")
    p = add_after(p, "除了“形状是否接近”和“性能是否接近”，本次新增图表还直接比较各模型相对 NN 的预测优劣。判断规则是：若某模型的 MSE vs y 小于 NN_mse_vs_y，则该 run 中该模型胜过 NN。nn_baseline 图左侧给出胜率，右侧给出相对 NN 的中位改善百分比；50% 和 0% 分别是胜负分界线。", "Normal")
    rows = baseline_rows("ALL")
    table = add_table_after(p, 4, 4)
    fill_table(table, ["模型", "有效 run", "相对 NN 胜率", "中位改善"], [
        [r["model"], str(int(r["n"])), pct(r["win_rate_vs_NN"]), imp(r["median_improve_vs_NN_pct"])]
        for _, r in rows.iterrows()
    ])
    p = add_after_element(table._element, p._parent, "总体看，FullPR 是唯一真正能和 NN 正面对打的模型，但全体合并后胜率为 38.9%、中位改善为 -85.3%，说明 NN 更稳。Taylor-PR 胜率为 0%，LayerTaylor-PR 总体胜率仅 5.6%，二者的主要价值是解释和复刻 NN，而不是击败 NN。", "Normal")
    p = add_caption_after(p, "图：paper_poly3 的 NN 基准胜负图。FullPR 多数时候胜过 NN。", "results_clean_paper_poly3_nn_baseline.png")
    p = add_caption_after(p, "图：paper_poly4 的 NN 基准胜负图。NN 在该欠设定高阶任务中更稳。", "results_clean_paper_poly4_nn_baseline.png")
    p = add_caption_after(p, "图：smooth_nonlinear 的 NN 基准胜负图。FullPR 略优于 NN。", "results_clean_smooth_nonlinear_nn_baseline.png")
    p = add_caption_after(p, "图：smooth_nonlinear_rand 的 NN 基准胜负图。NN 优势最明显。", "results_clean_smooth_nonlinear_rand_nn_baseline.png")
    p = add_after(p, "分数据集结论：paper_poly3 和 smooth_nonlinear 中，FullPR 更常胜过 NN；paper_poly4 和 smooth_nonlinear_rand 中，NN 更稳。这说明“神经网络可被多项式解释”与“多项式一定更准”是两个不同命题。几何等价性描述的是形状和性能差距之间的规律，预测胜负还取决于真函数阶数、随机非线性、特征族是否匹配以及训练稳定性。", "Normal")

    set_text(find_para(doc, contains="几何等价性成立且普适"),
             "几何等价性成立且普适：神经网络与多项式回归在形状上高度接近，且接近度强烈预测性能接近度（清洗后全体 ρ=0.839），跨命中/欠设定/非多项式数据普遍成立。这把“NN 是否在做多项式回归”从模糊直觉变成可测量、可证伪的结论。")
    add_bullets_after(find_para(doc, contains="方法论警示"),
                      [
        "以 NN 为基准的胜负结论更细：FullPR 是真正的预测竞争者，但只在 paper_poly3 与 smooth_nonlinear 中多数胜出；在 paper_poly4 与 smooth_nonlinear_rand 中 NN 更稳。Taylor 类方法不能被理解为稳定优于 NN 的预测模型，而应定位为解释/复刻工具。",
    ], "List Number")

    out = ROOT / "大报告_中文_完整版_改写版.docx"
    doc.save(str(out))
    return out


def english_report():
    doc = Document(str(ROOT / "big_report_en_complete.docx"))

    set_text(find_para(doc, contains="This report integrates: mathematical foundations"),
             "This report integrates: mathematical foundations, methodology, a panel-by-panel reading of all experimental figures, the quantitative correlation results from 3000 repeated simulations, the NN-baseline win/loss analysis, and conclusions ready for a paper.")
    set_text(find_para(doc, contains="This study answers a fundamental question"),
             "This study answers a fundamental question: are neural networks secretly doing polynomial regression? We use shape distance as a geometric ruler to turn this vague intuition into a measurable scientific claim, tested across 5 datasets, 3 activation functions, 3 network depths, and 3000 repeated simulations. The new NN-baseline figures further answer which model is actually more accurate: not every polynomial beats the NN; FullPR wins only in some cleaner structural settings. The core findings:")
    set_text(doc.paragraphs[10],
             "The core claim holds universally: the geometric closeness between a neural network and polynomial regression strongly predicts their performance closeness (cleaned overall Spearman rho = 0.839, p ~ 0; per-dataset rho ranges 0.440-0.870, all significant).")
    set_text(doc.paragraphs[13],
             "LayerTaylorPR fails rarely but catastrophically: in most cases it replicates the NN, but once the synaptic potential leaves the convergence region, MSE can explode. In the NN-baseline prediction comparison, its overall win rate is only 5.6%, so it is best understood as an explanation/replication tool rather than a robust predictive substitute.")
    add_bullets_after(doc.paragraphs[13], [
        "The NN-baseline comparison is not one-sided: FullPR wins more often in paper_poly3 (59.5% win rate, +62.6% median improvement) and smooth_nonlinear (58.5%, +37.0%), while the NN is more stable in paper_poly4 (FullPR win rate 36.7%, -97.0% median improvement) and smooth_nonlinear_rand (14.4%, -169.1%).",
    ], "List Number")

    set_text(find_para(doc, contains="Variants: NN_mse_vs_y"),
             "Variants: NN_mse_vs_y (NN vs truth), FPR_mse_vs_y (FullPR vs truth), FPR_mse_vs_NN (FullPR vs NN output), and LTPR_mse_vs_NN (LTPR vs NN output). The last two use the NN output as truth because they ask how closely a model replicates the NN. The new NN-baseline variables are: *_mse_delta_vs_NN (model MSE minus NN MSE; negative means the model is more accurate), *_improve_vs_NN_pct (relative improvement over NN; positive means better than NN), and *_better_than_NN (whether the model beats NN in that run).")
    set_text(find_para(doc, contains="Panel 4.  NN MSE vs FullPR MSE"),
             "Panel 4.  NN MSE vs FullPR MSE (y=x line): points below the line mean FullPR is more accurate; points above mean the NN is more accurate. The new nn_baseline figure aggregates this visual comparison into win rate and median improvement, which is the clearer way to report which model is better.")

    anchor = find_para(doc, exact="Part 4  Quantitative Findings")
    p = add_after(anchor, "4.0  Model Wins and Losses Against the NN Baseline", "Heading 2")
    p = add_after(p, "Beyond asking whether shapes are close and whether performance gaps are small, the new figures directly compare each model against the NN as the prediction baseline. The rule is simple: if a model's MSE vs y is smaller than NN_mse_vs_y, that model wins the run. The left panel of each nn_baseline figure reports win rate; the right panel reports median percentage improvement over NN. The 50% and 0% reference lines are the win/loss boundaries.", "Normal")
    rows = baseline_rows("ALL")
    table = add_table_after(p, 4, 4)
    fill_table(table, ["Model", "Valid runs", "Win rate vs NN", "Median improvement"], [
        [r["model"], str(int(r["n"])), pct(r["win_rate_vs_NN"]), imp(r["median_improve_vs_NN_pct"])]
        for _, r in rows.iterrows()
    ])
    p = add_after_element(table._element, p._parent, "Overall, FullPR is the only model that genuinely competes with the NN, but pooled across all datasets it has a 38.9% win rate and a -85.3% median improvement, so the NN is more stable overall. Taylor-PR has a 0% win rate, and LayerTaylor-PR wins only 5.6% of runs; these methods should be treated as explanatory or replicative devices rather than models that reliably outperform the NN.", "Normal")
    p = add_caption_after(p, "Figure: NN-baseline comparison for paper_poly3. FullPR wins more often than the NN.", "results_clean_paper_poly3_nn_baseline.png")
    p = add_caption_after(p, "Figure: NN-baseline comparison for paper_poly4. The NN is more stable in this underspecified higher-order task.", "results_clean_paper_poly4_nn_baseline.png")
    p = add_caption_after(p, "Figure: NN-baseline comparison for smooth_nonlinear. FullPR slightly outperforms the NN.", "results_clean_smooth_nonlinear_nn_baseline.png")
    p = add_caption_after(p, "Figure: NN-baseline comparison for smooth_nonlinear_rand. The NN has the clearest advantage.", "results_clean_smooth_nonlinear_rand_nn_baseline.png")
    p = add_after(p, "Dataset-level conclusion: FullPR wins more often in paper_poly3 and smooth_nonlinear, while the NN is more stable in paper_poly4 and smooth_nonlinear_rand. Thus, 'the NN can be explained by a polynomial' and 'a polynomial is always more accurate than the NN' are different claims. Geometric equivalence describes the relationship between shape and performance gap; the actual winner depends on polynomial order, random nonlinearity, feature specification, and training stability.", "Normal")

    set_text(find_para(doc, contains="Geometric equivalence holds universally"),
             "Geometric equivalence holds universally: neural networks are highly close to polynomials in shape, and closeness strongly predicts performance closeness (cleaned overall rho=0.839), holding across well-specified, underspecified, and non-polynomial data. This turns 'is the NN doing polynomial regression' from a vague intuition into a measurable, falsifiable conclusion.")
    add_bullets_after(find_para(doc, contains="Methodological caution"),
                      [
        "The NN-baseline win/loss result is more nuanced: FullPR is the true predictive competitor, but it wins mainly in paper_poly3 and smooth_nonlinear; in paper_poly4 and smooth_nonlinear_rand the NN is more stable. Taylor-style methods should not be interpreted as prediction models that reliably beat the NN; their main role is explanation and replication.",
    ], "List Number")

    out = ROOT / "big_report_en_complete_rewritten.docx"
    doc.save(str(out))
    return out


if __name__ == "__main__":
    print(chinese_report())
    print(english_report())
