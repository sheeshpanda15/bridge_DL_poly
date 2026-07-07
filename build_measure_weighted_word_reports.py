from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent

SUMMARY = ROOT / "measure_weighted_grand_mean_strategy_summary.csv"
CASE_SUMMARY = ROOT / "measure_weighted_grand_mean_case_summary.csv"
PAIRWISE = ROOT / "measure_weighted_grand_mean_pairwise.csv"
DISTANCES = ROOT / "measure_weighted_grand_mean_distances_combined.csv"

FIG_FLOW = ROOT / "measure_weighted_sampling_flowchart_cn.png"
FIG_GAIN = ROOT / "measure_weighted_grand_gain_by_p.png"
FIG_LEARNING = ROOT / "measure_weighted_grand_learning_gain.png"
FIG_CASE = ROOT / "measure_weighted_grand_case_gain_heatmap.png"
FIG_WEIGHT = ROOT / "measure_weighted_grand_weight_heatmap.png"

OUT_CN = ROOT / "大报告_中文_完整版_测度加权大实验更新版.docx"
OUT_EN = ROOT / "big_report_en_complete_measure_weighted_grand_update.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x0B, 0x25, 0x45)
MUTED = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def set_run_font(run, size=None, bold=None, color=None, east_asia="Microsoft YaHei"):
    run.font.name = "Calibri"
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_doc(doc, lang):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run(
        "Measure-weighted PR transfer experiment" if lang == "en" else "距离测度加权采样大实验"
    )
    set_run_font(r, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Generated from reproducible experiment outputs")
    set_run_font(r, size=9, color=MUTED)


def para(doc, text="", style=None, size=None, bold=None, color=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_title(doc, title, subtitle, meta):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=22, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(subtitle)
    set_run_font(r, size=12, color=MUTED)

    table = doc.add_table(rows=0, cols=2)
    widths = [2100, 7260]
    for label, value in meta:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
    set_table_geometry(table, widths)
    for row in table.rows:
        set_cell_shading(row.cells[0], HEADER_FILL)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_run_font(r, size=9)
        for r in row.cells[0].paragraphs[0].runs:
            r.font.bold = True
            r.font.color.rgb = DARK_BLUE
    doc.add_paragraph()


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + " ")
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        set_cell_shading(cell, HEADER_FILL)
    set_repeat_table_header(table.rows[0])
    for row_vals in rows:
        row = table.add_row()
        for i, value in enumerate(row_vals):
            row.cells[i].text = str(value)
    set_table_geometry(table, widths)
    for ridx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_run_font(r, size=8.7, bold=(ridx == 0), color=(DARK_BLUE if ridx == 0 else None))
    doc.add_paragraph()
    return table


def add_picture(doc, path, caption, width=6.3):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=9, color=MUTED)


def fmt_pct(x):
    return f"{x:.2f}%"


def fmt_rate(x):
    return f"{100 * x:.1f}%"


def load_data():
    return {
        "summary": pd.read_csv(SUMMARY),
        "case_summary": pd.read_csv(CASE_SUMMARY),
        "pairwise": pd.read_csv(PAIRWISE),
        "distances": pd.read_csv(DISTANCES),
    }


def strategy_rows(summary):
    rows = []
    for _, r in summary.iterrows():
        rows.append(
            [
                int(r["p_dim"]),
                r["strategy_label"],
                f"{r['final_mse_mean']:.6f}",
                fmt_pct(r["gain_vs_random_mean"]),
                fmt_rate(r["positive_gain_rate"]),
                int(r["runs"]),
            ]
        )
    return rows


def case_rows(case_summary, lang):
    mw = case_summary[case_summary["strategy"] == "measure_weighted"].copy()
    rows = []
    for _, r in mw.iterrows():
        label = r["case_label"]
        if lang == "cn":
            label = {
                "Quadratic": "二次多项式",
                "Smooth": "平滑非线性",
                "Strong nonlinear": "强非线性",
                "Local interior": "局部内部结构",
                "High-frequency": "高频非多项式",
            }.get(label, label)
        rows.append(
            [
                int(r["p_dim"]),
                label,
                f"{r['final_mse_mean']:.6f}",
                fmt_pct(r["gain_vs_random_mean"]),
                fmt_rate(r["positive_gain_rate"]),
            ]
        )
    return rows


def weight_rows(distances, lang):
    grouped = (
        distances.groupby(["p_dim", "case_label"], as_index=False)
        .agg(
            fpr=("fpr_distance", "mean"),
            tpr=("tpr_distance", "mean"),
            weight=("dopt_weight", "mean"),
        )
        .sort_values(["p_dim", "case_label"])
    )
    rows = []
    for _, r in grouped.iterrows():
        label = r["case_label"]
        if lang == "cn":
            label = {
                "Quadratic": "二次多项式",
                "Smooth": "平滑非线性",
                "Strong nonlinear": "强非线性",
                "Local interior": "局部内部结构",
                "High-frequency": "高频非多项式",
            }.get(label, label)
        rows.append([int(r["p_dim"]), label, f"{r['fpr']:.5f}", f"{r['tpr']:.5f}", f"{r['weight']:.3f}"])
    return rows


def build_cn(data):
    doc = Document()
    style_doc(doc, "cn")
    add_title(
        doc,
        "距离测度加权采样的大实验更新版",
        "从 NN-PR/TYPR 距离出发，决定 D-optimal 与随机采样的批次混合比例",
        [
            ("实验规模", "n=10000；p=20 和 p=50；5 类数据集；5 个 data seed；2 个初始采样 seed"),
            ("策略数量", "Measure-weighted、Random、D-optimal、Latin hypercube 共 4 组"),
            ("批次设置", "初始 5% 均匀 pilot；每轮新增 500 点；6 轮 batch 更新"),
            ("主规则", "使用 mean(d_FPR, d_TYPR) 作为保守距离，映射到 30%-95% 的 D-optimal 比例"),
        ],
    )
    add_callout(
        doc,
        "核心结论",
        "新的大实验表明，该测度最适合被解释为 transfer gate：距离小则更大胆迁移 PR/D-optimal 技术，距离大则保留更多随机探索。它不是万能地击败所有采样规则，但能揭示何时旧技术会产生负迁移。",
    )

    doc.add_heading("1. 实验流程", level=1)
    para(
        doc,
        "实际应用中无法预先知道 NN 与 FullPR/TYPR 是否接近。因此实验先在训练集上做均匀 pilot 抽样，训练初始 NN，再计算该 NN 与 FullPR、TYPR 的响应曲面距离。距离经过归一化后转成下一批数据中的 D-optimal 比例。",
    )
    add_picture(doc, FIG_FLOW, "图 1. 距离测度加权的批次采样流程", width=6.4)

    doc.add_heading("2. 实验设计", level=1)
    add_numbered(
        doc,
        [
            "对每个数据集做 75/25 train/test 拆分，训练集内先均匀抽取 5% 作为 pilot subset。",
            "用 pilot subset 训练单隐层 tanh NN；在 pilot validation 上计算 NN-FPR 和 NN-TYPR 距离。",
            "主实验使用 mean(d_FPR, d_TYPR) 而不是 min(...)，避免高维下单个 PR 代理偶然贴近导致过度使用 D-optimal。",
            "每轮固定新增 500 点，Measure-weighted 策略按 weight 混合 D-optimal 点和随机点；对照组为全随机、全 D-optimal 和 Latin hypercube。",
        ],
    )

    doc.add_heading("3. 主结果", level=1)
    add_table(
        doc,
        ["p", "策略", "Final MSE", "相对随机提升", "正提升率", "运行数"],
        strategy_rows(data["summary"]),
        [700, 2600, 1600, 1700, 1500, 1260],
    )
    add_picture(doc, FIG_GAIN, "图 2. 不同维度下各策略相对随机采样的最终提升", width=6.3)
    add_picture(doc, FIG_LEARNING, "图 3. 多轮 batch 更新中的相对随机提升曲线", width=6.3)

    doc.add_heading("4. 分数据集结果", level=1)
    para(
        doc,
        "Measure-weighted 在 p=20 的多数结构上为正提升，在 p=50 中能明显缓解全 D-optimal 的负迁移，尤其在 high-frequency 中将全 D-optimal 的大幅损失压回接近随机水平，在 local interior 中则取得最强结果。",
    )
    add_table(
        doc,
        ["p", "数据集", "Measure-weighted MSE", "相对随机提升", "正提升率"],
        case_rows(data["case_summary"], "cn"),
        [700, 2400, 2200, 2000, 1700],
    )
    add_picture(doc, FIG_CASE, "图 4. Measure-weighted 在各数据集和维度下的相对随机提升", width=5.9)

    doc.add_heading("5. 权重校准与测度解释", level=1)
    para(
        doc,
        "权重热图显示，采样比例不是按数据集标签硬编码，而是由每个 pilot run 的距离决定。p=20 中 high-frequency 被压到低 D-optimal 比例；p=50 中某些情形的权重明显下降，说明维度升高后 PR/TYPR 近似性变弱时，策略会自动更保守。",
    )
    add_table(
        doc,
        ["p", "数据集", "FPR 距离均值", "TYPR 距离均值", "D-optimal 比例均值"],
        weight_rows(data["distances"], "cn"),
        [650, 2200, 1900, 1900, 2310],
    )
    add_picture(doc, FIG_WEIGHT, "图 5. 由 NN-PR/TYPR 距离学习到的 D-optimal 平均比例", width=5.9)

    doc.add_heading("6. 敏感性与边界条件", level=1)
    add_callout(
        doc,
        "重要修正",
        "早先使用 min(d_FPR, d_TYPR) 时，p=50 的 Measure-weighted 相对随机为 -2.821%；改用 mean(d_FPR, d_TYPR) 后提高到 +0.501%，并超过全 D-optimal 的 -1.846%。这说明高维下需要更保守的距离组合规则。",
    )
    add_bullets(
        doc,
        [
            "优势：该测度能把采样策略从固定规则变成数据驱动规则，尤其能识别 D-optimal 可能负迁移的场景。",
            "限制：Latin hypercube 在某些高频或全局覆盖需求强的场景仍然更稳定；当前策略只在随机和 D-optimal 之间混合，尚未把 Latin 作为可选动作。",
            "后续改进：可把 D-optimal 下限从 30% 降到 10% 或引入三路混合（random / D-optimal / Latin），让远距离场景更接近纯探索。",
        ],
    )

    doc.add_heading("7. 可写入论文的表述", level=1)
    para(
        doc,
        "本实验支持的不是“测度加权采样永远优于所有设计”，而是一个更稳健的结论：NN-PR/TYPR 距离可以作为旧 PR/D-optimal 技术迁移到新 NN 模型前的可操作 gate。距离小，旧技术可高比例迁移；距离大，策略保留随机探索以降低负迁移风险。",
    )
    doc.save(OUT_CN)


def build_en(data):
    doc = Document()
    style_doc(doc, "en")
    add_title(
        doc,
        "Measure-Weighted Sampling: Grand Experiment Update",
        "Using NN-PR/Taylor-PR distance as an operational gate for D-optimal transfer",
        [
            ("Scale", "n=10000; p=20 and p=50; five dataset families; five data seeds; two pilot seeds"),
            ("Strategies", "Measure-weighted, random, D-optimal, and Latin hypercube"),
            ("Batch design", "5% uniform pilot, 500 new points per round, six batch updates"),
            ("Main rule", "mean(d_FPR, d_TYPR) mapped to a 30%-95% D-optimal fraction"),
        ],
    )
    add_callout(
        doc,
        "Main finding",
        "The distance measure is best interpreted as a transfer gate. Small NN-PR/Taylor-PR distance justifies more D-optimal transfer; large distance preserves random exploration and reduces negative transfer risk.",
    )

    doc.add_heading("1. Workflow", level=1)
    para(
        doc,
        "In real applications the distance between a trained NN and a PR model is not known in advance. The workflow therefore starts with a uniform pilot sample, trains an initial NN, measures NN-FPR and NN-Taylor-PR distances, normalizes the distance, and uses the resulting weight to choose the D-optimal/random mix in the next batch.",
    )
    add_picture(doc, FIG_FLOW, "Figure 1. Measure-weighted batch sampling workflow", width=6.4)

    doc.add_heading("2. Experimental Design", level=1)
    add_numbered(
        doc,
        [
            "Each dataset is split into 75% training and 25% testing; 5% of the training split is used as a uniform pilot subset.",
            "A single-hidden-layer tanh NN is trained on the pilot subset; NN-FPR and NN-Taylor-PR distances are measured on pilot validation points.",
            "The main rule uses mean(d_FPR, d_TYPR), not min(...), so high-dimensional false positives do not trigger excessive D-optimal transfer.",
            "Each batch adds 500 points. The measure-weighted policy mixes D-optimal and random points; controls use all-random, all-D-optimal, and Latin hypercube selection.",
        ],
    )

    doc.add_heading("3. Main Results", level=1)
    add_table(
        doc,
        ["p", "Strategy", "Final MSE", "Gain vs random", "Positive rate", "Runs"],
        strategy_rows(data["summary"]),
        [700, 2600, 1600, 1700, 1500, 1260],
    )
    add_picture(doc, FIG_GAIN, "Figure 2. Final gain versus random sampling by dimension", width=6.3)
    add_picture(doc, FIG_LEARNING, "Figure 3. Learning-curve gain over multiple batch rounds", width=6.3)

    doc.add_heading("4. Case-Level Results", level=1)
    para(
        doc,
        "The measure-weighted policy is positive on most p=20 cases and mitigates the p=50 failure of all-D-optimal sampling. It is strongest on local-interior structure and remains useful in high-frequency stress tests, although Latin hypercube can still be stronger when global coverage dominates.",
    )
    add_table(
        doc,
        ["p", "Dataset", "Measure-weighted MSE", "Gain vs random", "Positive rate"],
        case_rows(data["case_summary"], "en"),
        [700, 2400, 2200, 2000, 1700],
    )
    add_picture(doc, FIG_CASE, "Figure 4. Case-level gain of the measure-weighted policy", width=5.9)

    doc.add_heading("5. Weight Calibration", level=1)
    para(
        doc,
        "The learned weights are not hard-coded by dataset label. They change with the pilot run. This is the operational value of the measure: it converts an observed NN-PR/Taylor-PR relation into a concrete sampling fraction.",
    )
    add_table(
        doc,
        ["p", "Dataset", "Mean FPR distance", "Mean TYPR distance", "Mean D-opt fraction"],
        weight_rows(data["distances"], "en"),
        [650, 2200, 1900, 1900, 2310],
    )
    add_picture(doc, FIG_WEIGHT, "Figure 5. Mean D-optimal fraction learned from the distance measure", width=5.9)

    doc.add_heading("6. Sensitivity and Boundary Conditions", level=1)
    add_callout(
        doc,
        "Important correction",
        "The earlier min(d_FPR, d_TYPR) rule produced -2.821% relative gain for p=50. The conservative mean-distance rule improves this to +0.501% and beats all-D-optimal, which remains at -1.846%.",
    )
    add_bullets(
        doc,
        [
            "Strength: the measure turns sampling from a fixed rule into a data-driven transfer decision.",
            "Limitation: Latin hypercube remains competitive in high-frequency/global-coverage cases; the current policy mixes only random and D-optimal points.",
            "Next step: introduce a three-way policy over random, D-optimal, and Latin sampling or lower the D-optimal floor in far-distance cases.",
        ],
    )

    doc.add_heading("7. Paper-Ready Interpretation", level=1)
    para(
        doc,
        "The experiment does not support the claim that measure-weighted sampling is universally superior to every design. It supports a more defensible claim: NN-PR/Taylor-PR distance is an actionable transfer gate for deciding when legacy PR/D-optimal technology can be used safely with a new NN model.",
    )
    doc.save(OUT_EN)


def main():
    data = load_data()
    build_cn(data)
    build_en(data)
    print(f"Wrote {OUT_CN}")
    print(f"Wrote {OUT_EN}")


if __name__ == "__main__":
    main()
